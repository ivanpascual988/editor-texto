"""
Editor de texto que simula a cualquier editor de texto

@Author Iván Pascual
"""

# Importo las bibliotecas necesarias
import os
import tkinter as tk
from tkinter import scrolledtext
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
import docx


class TextEditor(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Editor de texto by Ivan Pascual") # Titulo de la app
        self.window_definition()
        

    def window_definition(self):
        """
        Defino la ventana al completo de la app
        """
        # Estilo del boton negrita
        n = ttk.Style()
        n.configure("Negrita.TButton", foreground="black", font=("Arial", 10, "bold"))
        # Botón poner en negrita el texto
        self.negrita = ttk.Button(self, text="N", width=5, style="Negrita.TButton", command=self.poner_quitar_negrita)
        self.negrita.place(x=50, y=10)

        # Estilo del boton cursiva
        k = ttk.Style()
        k.configure("Cursiva.TButton", foreground="black", font=("Arial", 10, "italic"))
        # Botón para poner en cursiva el texto
        self.cursiva = ttk.Button(self, text="K", width=5, style="Cursiva.TButton", command=self.poner_quitar_cursiva)
        self.cursiva.place(x=100, y=10)

        # Estilo del boton subrayado
        s = ttk.Style()
        s.configure("Subrayado.TButton", foreground="black", font=("Arial", 10))
        # Botón para subrayar el texto
        self.subrayado = ttk.Button(self, text="S", width=5, style="Subrayado.TButton", command=self.poner_quitar_subrayado)
        self.subrayado.configure(underline=False) # Añadimos el subrayado a la "S"
        self.subrayado.place(x=150, y=10)

        # Estilo del boton tachar
        t = ttk.Style()
        t.configure("Tachado.TButton", foreground="black", font=("Arial", 10, "overstrike"))
        # Botón para tachar el texto
        self.tachado = ttk.Button(self, text="abc", width=5, style="Tachado.TButton", command=self.poner_quitar_tachado)
        self.tachado.place(x=200, y=10)

        # Botón poner abrir un archivo de texto
        self.abrir = ttk.Button(self, text="Abrir archivo", command=self.abrir_archivo)
        self.abrir.place(x=420, y=10)

        # Botón guardar el texto escrito
        self.guardar = ttk.Button(self, text="Guardar archivo", command=self.guardar_archivo)
        self.guardar.place(x=520, y=10)

        # Área de entrada de texto para el usuario
        self.texto = scrolledtext.ScrolledText(self, wrap=tk.WORD, height=20, width=80, font=("Arial", 10))
        self.texto.pack(pady=50, padx=50)
    
    def poner_quitar_negrita(self):
        """
        Funcion para comprobar si el texto seleccionado esta en negrita o no.
        Si esta en negrita lo quita y si no esta en negrita, lo pone en negrita
        """
        # Obtengo el texto seleccionado
        inicio_texto = self.texto.index(tk.SEL_FIRST)
        fin_texto =  self.texto.index(tk.SEL_LAST)

        # Obtengo las etiquetas del texto seleccionado
        etiqueta_texto = (self.texto.tag_names(inicio_texto))

        # Si tiene la etiqueta bold la eliminamos
        if "bold" in etiqueta_texto:
            self.texto.tag_remove("bold", inicio_texto, fin_texto)
            if "italic" in etiqueta_texto:
                self.texto.tag_configure("style", font=("Arial", 10, "normal", "italic"))
            else:
                self.texto.tag_remove("style", inicio_texto, fin_texto)
        # Si no tiene la etiqueta la añadimos
        else:
            self.texto.tag_add("bold", inicio_texto, fin_texto) 
            # Si no tiene la etiqueta style la añadimos
            if "style" not in etiqueta_texto:
                self.texto.tag_add("style", inicio_texto, fin_texto)
            # Si tiene la etiqueta bold
            if "italic" in etiqueta_texto:
                self.texto.tag_configure("style", font=("Arial", 10, "bold", "italic"))
            else:
                self.texto.tag_configure("style", font=("Arial", 10, "bold", "roman"))
        
    
    def poner_quitar_cursiva(self):
        """
        Funcion para comprobar si el texto seleccionado esta en cursiva o no.
        Si esta en cursiva lo quita y si no esta en cursiva, lo pone en cursiva
        """
        # Obtengo el texto seleccionado
        inicio_texto = self.texto.index(tk.SEL_FIRST)
        fin_texto =  self.texto.index(tk.SEL_LAST)

        # Obtengo las etiquetas del texto seleccionado
        etiqueta_texto = (self.texto.tag_names(inicio_texto))

        # Si tiene la etiqueta italic la eliminamos
        if "italic" in etiqueta_texto:
            self.texto.tag_remove("italic", inicio_texto, fin_texto)
            if "bold" in etiqueta_texto:
                self.texto.tag_configure("style", font=("Arial", 10, "bold", "roman"))
            else:
                self.texto.tag_remove("style", inicio_texto, fin_texto)
        # Si no tiene la etiqueta la añadimos
        else:
            self.texto.tag_add("italic", inicio_texto, fin_texto) 
            # Si no tiene la etiqueta style la añadimos
            if "style" not in etiqueta_texto:
                self.texto.tag_add("style", inicio_texto, fin_texto)
            # Si tiene la etiqueta bold
            if "bold" in etiqueta_texto:
                self.texto.tag_configure("style", font=("Arial", 10, "bold", "italic"))
            else:
                self.texto.tag_configure("style", font=("Arial", 10, "normal", "italic"))
            
        
    def poner_quitar_subrayado(self):
        """
        Funcion para comprobar si el texto seleccionado esta subrayado o no.
        Si esta subrayado lo quita y si no esta subrayado, lo subraya
        """
        # Obtengo el texto seleccionado
        inicio_texto = self.texto.index(tk.SEL_FIRST)
        fin_texto =  self.texto.index(tk.SEL_LAST)

        # Obtenemos las etiquetas del texto seleccionado
        etiqueta_texto = self.texto.tag_names(inicio_texto)

        # Si tiene la etiqueta se la eliminamos
        if "underline" in etiqueta_texto:
            self.texto.tag_remove("underline", inicio_texto, fin_texto)
        # Si no tiene la etiqueta la añadimos
        else:
            self.texto.tag_add("underline", inicio_texto, fin_texto)
            self.texto.tag_configure("underline", underline=True)   # Configuramos la etiqueta para que esta subrayado

    def poner_quitar_tachado(self):
        """
        Funcion para comprobar si el texto seleccionado esta tachado o no.
        Si esta tachado lo quita y si no esta tachado, lo tacha
        """
        # Obtengo el texto seleccionado
        inicio_texto = self.texto.index(tk.SEL_FIRST)
        fin_texto =  self.texto.index(tk.SEL_LAST)

        # Compruebo si esta subrayado o no
        etiqueta_texto = self.texto.tag_names(inicio_texto)

        # Si tiene la etiqueta se la eliminamos
        if "overstrike" in etiqueta_texto:
            self.texto.tag_remove("overstrike", inicio_texto, fin_texto)
        # Si no tiene la etiqueta la añadimos
        else:
            self.texto.tag_add("overstrike", inicio_texto, fin_texto)
            self.texto.tag_configure("overstrike", overstrike=True)    # Configuramos la etiqueta para que esta tachado
        
    def abrir_archivo(self):
        """
        Funcion para cargar un archivo de texto ya existente en el recuadro del editor
        """
        # Guardamos la ruta del archivo seleccionado, unicamente se pondran subir archivos de texto y de word
        ruta_archivo= filedialog.askopenfilename(
            filetypes=(
                ("Archivos de texto", "*.txt"),
                ("Archivos de Word", "*.docx")
            )
        )

        # Controlo que se suba un archivo
        try:
            archivo = open(ruta_archivo, "r", encoding="utf-8")     # Abro el archivo
            ruta, extension = os.path.splitext(archivo.name)        # Obtengo la extension

            # Si es un archivo de texto
            if extension == ".txt":
                contenido = archivo.read()
                self.texto.delete("1.0", tk.END)    # Borro el contenido que habia
                self.texto.insert("1.0", contenido) # Inserto el contenido
            
            # Si es un archivo de word
            if extension == ".docx":
                documento = docx.Document(archivo.name)
                self.texto.delete("1.0", tk.END)    # Borro el contendio que habia
                # Recorro el archivo word parrafo a parrafo
                for parrafo in documento.paragraphs:
                    # Inserto el contenido del parrago y paso a la siguiente linea
                    self.texto.insert("1.0", parrafo.text + "\n")
        
        # Mensaje si no se sube ningun archivo
        except FileNotFoundError:
            messagebox.showerror("Atencion", "No selecciono ningun archivo")
    
    def guardar_archivo(self):
        """
        Funcion para guardar el texto en un archivo
        """
         # Obtener el texto
        contenido = self.texto.get("1.0", tk.END) 
        ruta_guardado = filedialog.asksaveasfilename(
            defaultextension=".txt",    # extension por defecto
            filetypes=(
                ("Archivos de texto", "*.txt"),
                ("Archivos de Word", "*.docx")
            )
        )
    
        # Si guardamos el archivo
        if ruta_guardado:

            # Comprobamos que el archivo se guarda correctamente
            try:
                with open(ruta_guardado, "w", encoding="utf-8") as archivo:
                    ruta, extension = os.path.splitext(archivo.name)
                    # Si es un archivo de texto
                    if extension == ".txt":
                        archivo.write(contenido)
                        messagebox.showinfo("Información", "El archivo se ha guardado correctamente.")

                    # Si es un archivo de word
                    if extension == ".docx":
                        doc = docx.Document()
                        doc.add_paragraph(contenido)
                        doc.save(ruta_guardado)
                        messagebox.showinfo("Información", "El archivo se ha guardado correctamente.")

            except Exception as e:
                messagebox.showerror("Error", f"Error al guardar el archivo: {str(e)}")

if __name__ == "__main__":
    app = TextEditor()
    app.mainloop()